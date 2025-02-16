import csv
import docx
from django.core.files.base import File
import os
from pathlib import Path
from app.BunnyStorage import BunnyStorage
import uuid

class TestSeriesExtractor:
    def __init__(self, file: File, name: str, id):
        self.file = file
        self.name = name
        self.id = id

    def is_csv(self) -> bool:
        """Check if the file is a valid CSV"""
        if not self.file.name.lower().endswith(".csv"):
            return False
        return True

    def is_docx(self) -> bool:
        """Check if the file is a valid DOCX"""
        if not self.file.name.lower().endswith(".docx"):
            return False
        try:
            self.file.seek(0)  # Ensure reading from start
            docx.Document(self.file)
            return True
        except Exception:
            return False

    def extract_questions(self):
        """Extract questions based on file type"""
        if self.is_csv():
            return self.extract_from_csv()
        elif self.is_docx():
            return self.extract_from_docx()
        else:
            raise ValueError("Unsupported file format. Please provide a CSV or DOCX file.")

    def extract_from_csv(self):
        """Extract questions from a CSV file"""
        self.file.seek(0)  # Ensure reading from the start
        reader = csv.reader(self.file.read().decode("utf-8",errors="ignore").splitlines())
    
        headers = next(reader, None)  # Extract headers
        if not headers:
            return [], []  # Return empty lists if no headers exist

        structured_data = []
        solution_data = []

        for row in reader:
            if len(row) >= len(headers):
                question_data = {headers[i]: row[i] for i in range(len(headers))}

                question_uuid = str(uuid.uuid4())

                # Process options
                options = [question_data.get(f"option{chr(65+i)}", "") for i in range(4)]
                options = [opt for opt in options if opt]  # Remove empty options

                structured_entry = {
                    "UUID": question_uuid,
                    "Question": question_data.get("question", ""),
                    "Option": options,
                    "Type": question_data.get("type", ""),
                    "Negative Marks": question_data.get("negative_marks", ""),
                    "Positive Marks": question_data.get("positive_marks", "")
                }

                # Process correct option and explanation separately
                answer_solution = {
                    "UUID": question_uuid,
                    "Answer": question_data.get("correctOpt", ""),
                    "Solution": question_data.get("explanation", ""),
                    "Question": question_data.get("question", ""),
                    "Negative Marks": question_data.get("negative_marks", ""),
                    "Positive Marks": question_data.get("positive_marks", ""),
                }

                structured_data.append(structured_entry)
                solution_data.append(answer_solution)
                    
        return structured_data, solution_data
    
    
    def save_image(self, image_data, question_number, image_index,key):
        """Save image in structured format: file_folder/Q_NUMBER/solution/IMG_INDEX.jpg"""
        folder_path = f"test_series/{self.id}/{question_number}/{key}"

        image_filename = f"_{image_index}.jpg"
        storage = BunnyStorage(folder_path)
        try:
            response= storage._save(image_filename,image_data)
            return response
        except Exception as e:
            raise ValueError(e)


    def extract_images_from_document(self, document):
        """Extract all embedded images and return a mapping of rel ID to image data"""
        image_mapping = {}
        for rel in document.part.rels:
            try:
                rel_target = document.part.rels[rel].target_part
                if "image" in rel_target.content_type:
                    image_mapping[rel] = rel_target.blob
            except ValueError:
                continue
        return image_mapping
    

    def extract_from_docx(self):
        """Extract questions from a DOCX file"""
        if not self.file:
            return [], []

        try:
            document = docx.Document(self.file)
        except Exception as e:
            print(f"Error reading document: {e}")
            return [], []

        table_data = []
        answer_solution_data = []
        question_number = 0

        image_mapping = self.extract_images_from_document(document)

        for table in document.tables:
            table_dict = {}
            answer_solution = {}
            image_index = 1
            question_uuid = str(uuid.uuid4()) 
    
            for row in table.rows:
                    
                cells = [cell for cell in row.cells]

                for rel in cells[1]._element.findall(".//a:blip", namespaces={"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}):
                        rid = rel.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rid in image_mapping:
                            image_path = self.save_image(image_mapping[rid], question_number, image_index, cells[0].text.strip())
                            cells[1].text += f' <img src="{image_path}" />'
                            image_index += 1


                if len(cells) >= 2:
                    key, value = cells[0].text.strip(), cells[1].text.strip()

                    # Store "Answer" and "Solution" separately
                    if key in ["Answer", "Solution"]:
                        answer_solution[key] = value
                    else:
                        if key in table_dict:
                            if isinstance(table_dict[key], list):
                                table_dict[key].append(value)
                            else:
                                table_dict[key] = [table_dict[key], value]
                        else:
                            table_dict[key] = value

                    if key in ["Positive Marks","Negative Marks","Question"]:
                        answer_solution[key] = value
                        question_number +=1 

            if table_dict:
                table_dict["UUID"] = question_uuid
                table_data.append(table_dict)

            if answer_solution:
                answer_solution["UUID"] = question_uuid
                answer_solution_data.append(answer_solution)

        return table_data, answer_solution_data
