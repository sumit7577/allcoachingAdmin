from django.db import models
from BunnyCDN.Storage import Storage
from BunnyCDN.CDN import CDN
from app.BunnyStorage import BunnyStorage
from django.utils import timezone
from app.Bunny import TusFileUploader
from docx import Document



class User(models.Model):
    id = models.BigAutoField(primary_key=True)
    name = models.CharField(max_length=150)
    username = models.CharField(unique=True, max_length=100, blank=True, null=True)
    email = models.CharField(unique=True, max_length=100)
    password = models.CharField(max_length=400)
    phone = models.CharField(unique=True, max_length=13)
    state = models.CharField(max_length=100, blank=True, null=True)
    pincode = models.PositiveBigIntegerField(blank=True, null=True)
    address = models.CharField(max_length=400, blank=True, null=True)
    date_joined = models.DateTimeField(default=timezone.now())
    date_updated = models.DateTimeField(default=timezone.now())
    is_active = models.BooleanField(default=True)
    image = models.ImageField(storage=BunnyStorage(), blank=True, null=True)
    is_institute = models.BooleanField(default=False)

    def createFileImage(self):
        """
        Generate the BunnyCDN directory path for the image dynamically.
        """
        return f"user/{self.id}/"

    def save(self, *args, **kwargs):
        """
        Override the save method to set the BunnyStorage dynamically.
        """
        if not self._state.adding:  # Ensures that 'id' is available
            directory = self.createFileImage()
            self.image.storage = BunnyStorage(directory)

        super().save(*args, **kwargs)

    def __str__(self):
        return f'{self.name}'
    

    class Meta:
        managed = True
        db_table = 'user'



class AuthToken(models.Model):
    key = models.CharField(primary_key=True, max_length=40)
    created = models.DateTimeField(default=timezone.now())
    user = models.OneToOneField(User, on_delete=models.CASCADE)

    def __str__(self):
        return self.key

    class Meta:
        managed = True
        db_table = 'auth_token'


class Banner(models.Model):
    id = models.BigAutoField(primary_key=True)
    title = models.CharField(max_length=150)
    image = models.ImageField(storage=BunnyStorage())
    date_created = models.DateTimeField(default=timezone.now())
    date_updated = models.DateTimeField(default=timezone.now())

    def createFileImage(self):
        """
        Generate the BunnyCDN directory path for the image dynamically.
        """
        return f"banner/{self.id}/"

    def save(self, *args, **kwargs):
        """
        Override the save method to set the BunnyStorage dynamically.
        """
        if not self._state.adding:  # Ensures that 'id' is available
            directory = self.createFileImage()
            self.image.storage = BunnyStorage(directory)

        super().save(*args, **kwargs)
    
    def __str__(self):
        return self.title

    class Meta:
        managed = True
        db_table = 'banner'


class Category(models.Model):
    id = models.BigAutoField(primary_key=True)
    name = models.CharField(max_length=150)
    icon = models.CharField(max_length=200, blank=True, null=True)
    date_created = models.DateTimeField(default=timezone.now())
    date_updated = models.DateTimeField(default=timezone.now())

    def __str__(self):
        return self.name

    class Meta:
        managed = True
        db_table = 'category'

class Institute(models.Model):
    id = models.BigAutoField(primary_key=True)
    name = models.CharField(max_length=100, blank=True, null=True)
    about = models.TextField(blank=True, null=True)
    category = models.ForeignKey(Category, on_delete=models.CASCADE)
    banners = models.ManyToManyField(Banner)
    director_name = models.CharField(max_length=150)
    user = models.OneToOneField(User,on_delete=models.CASCADE)
    date_created = models.DateTimeField(default=timezone.now())
    date_updated = models.DateTimeField(default=timezone.now())
    image = models.ImageField(storage=BunnyStorage(), null=True, blank=True)

    def createFileImage(self):
        """
        Generate the BunnyCDN directory path for the image dynamically.
        """
        return f"institute/{self.id}/"

    def save(self, *args, **kwargs):
        """
        Override the save method to set the BunnyStorage dynamically.
        """
        if not self._state.adding:  # Ensures that 'id' is available
            directory = self.createFileImage()
            self.image.storage = BunnyStorage(directory)

        super().save(*args, **kwargs)

    def __str__(self):
        return self.name

    class Meta:
        managed = True
        db_table = 'institute'


class Course(models.Model):
    id = models.BigAutoField(primary_key=True)
    name = models.CharField(max_length=150)
    institute = models.ForeignKey(Institute, on_delete=models.CASCADE)
    category = models.ForeignKey(Category, on_delete=models.CASCADE)
    banner = models.ForeignKey(Banner, on_delete=models.CASCADE)
    collection = models.JSONField(blank=True, null=True)
    description = models.TextField(blank=True, null=True)
    price = models.FloatField()
    created_at = models.DateTimeField(default=timezone.now())
    updated_at = models.DateTimeField(default=timezone.now())
    image = models.ImageField(storage=BunnyStorage(), null=True, blank=True)

    def createFileImage(self):
        """
        Generate the BunnyCDN directory path for the image dynamically.
        """
        return f"course/{self.id}/"

    def save(self, *args, **kwargs):
        """
        Override the save method to set the BunnyStorage dynamically.
        """

        if self.pk:
            previous = Course.objects.filter(pk=self.pk).values("image").first()
            if previous and previous["image"] != self.image.name:
                directory = self.createFileImage()
                self.image.storage = BunnyStorage(directory)
                
            if self.collection is None:
                name = f"{self.institute.name}-{self.name}"
                collection = TusFileUploader(instance=None)
                status,value =  collection.createCollection(name)
                if not status:
                    raise Exception(f"Failed to create collection: {str(value)}")
                self.collection = value

        if self._state.adding:
            name = f"{self.institute.name}-{self.name}"
            collection = TusFileUploader(instance=None)
            status,value =  collection.createCollection(name)
            if not status:
                raise Exception(f"Failed to create collection: {str(value)}")
            self.collection = value
            directory = self.createFileImage()
            self.image.storage = BunnyStorage(directory)
                

        super().save(*args, **kwargs)



    def __str__(self):
        return self.name

    class Meta:
        managed = True
        db_table = 'course'


class CourseVideos(models.Model):
    id = models.BigAutoField(primary_key=True)
    name = models.CharField(max_length=150)
    course = models.ForeignKey(Course, on_delete=models.CASCADE)
    description = models.TextField(blank=True, null=True)
    video = models.FileField(storage=TusFileUploader(instance=None), blank=True, null=True)
    metadata = models.JSONField(blank=True, null=True)
    created_at = models.DateTimeField(default=timezone.now())
    updated_at = models.DateTimeField(default=timezone.now())


    def save(self, *args, **kwargs):
        """
        Override the save method to set the BunnyStorage dynamically.
        """
        self.video.storage = TusFileUploader(instance=self)
        super().save(*args, **kwargs)

    def __str__(self):
        return self.name

    class Meta:
        managed = True
        db_table = 'course_videos'

def docxReader(instance,filename):
    docx = Document(instance.file)


class TestSeries(models.Model):
    id = models.BigAutoField(primary_key=True)
    name = models.CharField(max_length=300, blank=True, null=True)
    file = models.FileField(storage=BunnyStorage(), blank=True, null=True)
    course = models.ForeignKey(to=Course,blank=True, null=True,on_delete=models.CASCADE)
    description = models.TextField(blank=True, null=True)
    questions = models.JSONField()
    timer = models.BigIntegerField()
    created_at = models.DateTimeField(default=timezone.now())
    updated_at = models.DateTimeField(default=timezone.now())


    class Meta:
        managed = True
        db_table = 'test_series'

    def __str__(self):
        return self.name
    
    def createDocDir(self):
        """
        Generate the BunnyCDN directory path for the image dynamically.
        """
        return f"course/{self.course.id}/docs/"

    def save(self, *args, **kwargs):
        """
        Override the save method to set the BunnyStorage dynamically.
        """
        directory = self.createDocDir()
        self.file.storage = BunnyStorage(directory)

        """Extract Question from files"""
        document = Document(self.file.file)
        table_data = [] 
        answer_solution_data = []

        for table in document.tables:
            table_dict = {}
            answer_solution = {}

            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:  # Ensure there are at least two cells
                    key, value = cells[0], cells[1]

                    # Store "Answer" and "Solution" separately
                    if key in ["Answer", "Solution"]:
                        answer_solution[key] = value
                    else:
                        if key in table_dict:  # If key already exists, convert value to list
                            if isinstance(table_dict[key], list):
                                table_dict[key].append(value)
                            else:
                                table_dict[key] = [table_dict[key], value]
                        else:
                            table_dict[key] = value

                    if key in ["Question"]:
                        answer_solution[key] = value

            if table_dict:  # Append only if there's valid data
                table_data.append(table_dict)
            if answer_solution:  # Append Answer & Solution separately
                answer_solution_data.append(answer_solution)
        
        self.questions = table_data
        TestSeriesSolution.objects.create(test_series= self,solution=answer_solution_data).save()

        if self.name == "":
            self.name = self.file.name

        super().save(*args, **kwargs)


class TestSeriesAttempt(models.Model):
    id = models.BigAutoField(primary_key=True)
    test_series = models.ForeignKey(to=TestSeries,on_delete=models.CASCADE)
    user = models.ForeignKey(to=User,on_delete=models.CASCADE)
    result = models.JSONField()
    created_at = models.DateTimeField(default=timezone.now())
    updated_at = models.DateTimeField(default=timezone.now())

    class Meta:
        managed = True
        db_table = 'test_series_attempt'

    def __str__(self):
        return self.user.username


class TestSeriesSolution(models.Model):
    id = models.BigAutoField(primary_key=True)
    test_series = models.ForeignKey(to=TestSeries,on_delete=models.CASCADE)
    description = models.TextField(blank=True, null=True)
    solution = models.JSONField()
    created_at = models.DateTimeField(default=timezone.now())
    updated_at = models.DateTimeField(default=timezone.now())

    class Meta:
        managed = True
        db_table = 'test_series_solution'

    def __str__(self):
        return self.test_series.name


"""class InstituteBanners(models.Model):
    id = models.BigAutoField(primary_key=True)
    institute_id = models.BigIntegerField()
    banner_id = models.BigIntegerField()

    class Meta:
        managed = True
        db_table = 'institute_banners'"""




"""class TestRel(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE)
    course = models.ForeignKey(Course, on_delete=models.CASCADE)
    banners = models.ManyToManyField(Banner)


    def __str__(self):
        return f'{self.user.name}'"""